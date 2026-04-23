"""
Build a formatted cover letter .docx from structured content.

Usage:
    python scripts/build_cover_letter.py cover_letter_content.json output.docx
"""

import json
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, bold=False, size=11):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"


def add_paragraph(doc, text, bold=False, size=11, space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, bold=bold, size=size)
    return p


def add_competency_paragraph(doc, label, body, size=11):
    """Paragraph with a bold label followed by regular body text."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)

    label_run = p.add_run(f"{label}: ")
    set_font(label_run, bold=True, size=size)

    body_run = p.add_run(body)
    set_font(body_run, bold=False, size=size)
    return p


def build(content: dict, output_path: str):
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Header: NAME
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(content["name"].upper())
    name_run.bold = True
    name_run.font.size = Pt(14)
    name_run.font.name = "Calibri"

    # Contact line
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(12)
    contact_run = contact_p.add_run(content["contact_line"])
    set_font(contact_run, size=10)

    # Date
    if content.get("date"):
        add_paragraph(doc, content["date"], size=11, space_after=4)

    # Recipient
    if content.get("recipient"):
        add_paragraph(doc, content["recipient"], size=11, space_after=12)

    # Salutation
    add_paragraph(doc, content["salutation"], size=11, space_after=8)

    # Opening paragraph
    add_paragraph(doc, content["opening"], size=11, space_after=8)

    # Competency sections
    for section in content["competencies"]:
        add_competency_paragraph(doc, section["label"], section["body"])

    # Closing paragraph
    add_paragraph(doc, content["closing"], size=11, space_after=8)

    # Sign-off
    for line in content["signoff"]:
        add_paragraph(doc, line, size=11, space_before=0, space_after=2)

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python build_cover_letter.py <content.json> <output.docx>")
        sys.exit(1)

    with open(sys.argv[1], encoding="utf-8") as f:
        content = json.load(f)

    build(content, sys.argv[2])
